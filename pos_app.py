from flask import Flask, render_template, request, redirect, url_for, session, flash, send_file
from datetime import datetime
import os
import shutil
from docx import Document
import psycopg2
from psycopg2 import sql
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

app = Flask(__name__, template_folder='templates')
app.secret_key = os.getenv('SECRET_KEY', 'change_this_to_a_secure_random_value')

@app.context_processor
def inject_datetime():
    return {'datetime': datetime}

# === Database Configuration ===
def get_db_connection():
    if 'RENDER' in os.environ:
        # Render production (free tier)
        db_url = os.getenv('DATABASE_URL')
        return psycopg2.connect(db_url, sslmode='require')  # Required for free tier
    else:
        # Local development
        return psycopg2.connect(
            host=os.getenv('DB_HOST'),
            database=os.getenv('DB_NAME'),
            user=os.getenv('DB_USER'),
            password=os.getenv('DB_PASSWORD'),
            port=os.getenv('DB_PORT')
        )
# --------------------------
# Database Initialization
# --------------------------
def initialize_database():
    conn = get_db_connection()
    cur = conn.cursor()
    
    # Create tables if they don't exist
    cur.execute("""
        CREATE TABLE IF NOT EXISTS users (
            id SERIAL PRIMARY KEY,
            username VARCHAR(50) UNIQUE NOT NULL,
            password VARCHAR(100) NOT NULL,
            role VARCHAR(20) NOT NULL
        )
    """)
    
    cur.execute("""
        CREATE TABLE IF NOT EXISTS products (
            id SERIAL PRIMARY KEY,
            name VARCHAR(100) UNIQUE NOT NULL,
            buy_price DECIMAL(10, 2) NOT NULL,
            sell_price DECIMAL(10, 2) NOT NULL,
            stock INTEGER NOT NULL
        )
    """)
    
    cur.execute("""
        CREATE TABLE IF NOT EXISTS oils (
            id SERIAL PRIMARY KEY,
            name VARCHAR(100) UNIQUE NOT NULL,
            buy_price DECIMAL(10, 2) NOT NULL,
            sell_price DECIMAL(10, 2) NOT NULL,
            stock INTEGER NOT NULL
        )
    """)
    
    cur.execute("""
        CREATE TABLE IF NOT EXISTS wheels (
            id SERIAL PRIMARY KEY,
            name VARCHAR(100) UNIQUE NOT NULL,
            buy_price DECIMAL(10, 2) NOT NULL,
            sell_price DECIMAL(10, 2) NOT NULL,
            stock INTEGER NOT NULL
        )
    """)
    
    cur.execute("""
        CREATE TABLE IF NOT EXISTS sales (
            id SERIAL PRIMARY KEY,
            product VARCHAR(100) NOT NULL,
            price DECIMAL(10, 2) NOT NULL,
            quantity INTEGER NOT NULL,
            total DECIMAL(10, 2) NOT NULL,
            sale_date TIMESTAMP NOT NULL,
            receipt_id VARCHAR(50) NOT NULL
        )
    """)
    
    cur.execute("""
        CREATE TABLE IF NOT EXISTS credit_sales (
            id SERIAL PRIMARY KEY,
            customer_name VARCHAR(100) NOT NULL,
            product VARCHAR(100) NOT NULL,
            price DECIMAL(10, 2) NOT NULL,
            quantity INTEGER NOT NULL,
            total DECIMAL(10, 2) NOT NULL,
            sale_date TIMESTAMP NOT NULL,
            receipt_id VARCHAR(50) NOT NULL,
            month_year VARCHAR(7) NOT NULL
        )
    """)
    
    cur.execute("""
        CREATE TABLE IF NOT EXISTS medgulf_sales (
            id SERIAL PRIMARY KEY,
            customer_name VARCHAR(100) NOT NULL,
            product VARCHAR(100) NOT NULL,
            price DECIMAL(10, 2) NOT NULL,
            quantity INTEGER NOT NULL,
            total DECIMAL(10, 2) NOT NULL,
            sale_date TIMESTAMP NOT NULL,
            receipt_id VARCHAR(50) NOT NULL,
            month_year VARCHAR(7) NOT NULL
        )
    """)
    
    # Create indexes for better performance
    cur.execute("CREATE INDEX IF NOT EXISTS sales_date_idx ON sales(sale_date)")
    cur.execute("CREATE INDEX IF NOT EXISTS credit_sales_month_idx ON credit_sales(month_year)")
    cur.execute("CREATE INDEX IF NOT EXISTS medgulf_sales_month_idx ON medgulf_sales(month_year)")
    
    # Check if we need to add default users
    cur.execute("SELECT COUNT(*) FROM users")
    if cur.fetchone()[0] == 0:
        cur.execute(
            "INSERT INTO users (username, password, role) VALUES (%s, %s, %s)",
            ('admin', 'admin123', 'admin')
        )
        cur.execute(
            "INSERT INTO users (username, password, role) VALUES (%s, %s, %s)",
            ('cashier', '1234', 'cashier')
        )
    
    conn.commit()
    cur.close()
    conn.close()

# --------------------------
# Authentication
# --------------------------
def validate_login(username, password):
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute(
        "SELECT role FROM users WHERE username = %s AND password = %s",
        (username.lower(), password))
    result = cur.fetchone()
    cur.close()
    conn.close()
    return result[0] if result else None

# --------------------------
# Inventory Management
# --------------------------
def get_products():
    """Get products with buying price (admin only)"""
    conn = get_db_connection()
    cur = conn.cursor()
    
    if session.get('role') == 'admin':
        cur.execute("SELECT name, buy_price, sell_price, stock FROM products ORDER BY name")
    else:
        cur.execute("SELECT name, sell_price, stock FROM products ORDER BY name")
    
    products = []
    for row in cur.fetchall():
        if session.get('role') == 'admin':
            product = {
                'name': row[0],
                'buy_price': float(row[1]),
                'price': float(row[2]),
                'stock': int(row[3])
            }
        else:
            product = {
                'name': row[0],
                'price': float(row[1]),
                'stock': int(row[2])
            }
        products.append(product)
    
    cur.close()
    conn.close()
    return products

def get_oils():
    """Get oils with buying price (admin only)"""
    conn = get_db_connection()
    cur = conn.cursor()
    
    if session.get('role') == 'admin':
        cur.execute("SELECT name, buy_price, sell_price, stock FROM oils ORDER BY name")
    else:
        cur.execute("SELECT name, sell_price, stock FROM oils ORDER BY name")
    
    oils = []
    for row in cur.fetchall():
        if session.get('role') == 'admin':
            oil = {
                'name': row[0],
                'buy_price': float(row[1]),
                'price': float(row[2]),
                'stock': int(row[3])
            }
        else:
            oil = {
                'name': row[0],
                'price': float(row[1]),
                'stock': int(row[2])
            }
        oils.append(oil)
    
    cur.close()
    conn.close()
    return oils

def get_wheels():
    """Get wheels with buying price (admin only)"""
    conn = get_db_connection()
    cur = conn.cursor()
    
    if session.get('role') == 'admin':
        cur.execute("SELECT name, buy_price, sell_price, stock FROM wheels ORDER BY name")
    else:
        cur.execute("SELECT name, sell_price, stock FROM wheels ORDER BY name")
    
    wheels = []
    for row in cur.fetchall():
        if session.get('role') == 'admin':
            wheel = {
                'name': row[0],
                'buy_price': float(row[1]),
                'price': float(row[2]),
                'stock': int(row[3])
            }
        else:
            wheel = {
                'name': row[0],
                'price': float(row[1]),
                'stock': int(row[2])
            }
        wheels.append(wheel)
    
    cur.close()
    conn.close()
    return wheels

def get_stock(table_name, item_name):
    """Check current stock level"""
    conn = get_db_connection()
    cur = conn.cursor()
    
    query = sql.SQL("SELECT stock FROM {} WHERE name = %s").format(
        sql.Identifier(table_name))
    
    cur.execute(query, (item_name,))
    result = cur.fetchone()
    cur.close()
    conn.close()
    
    return result[0] if result else 0

def pending_in_session(item_name, item_type='product'):
    """Check for pending items in current session"""
    pending = 0
    items = session.get('receipt_items', [])
    for it in items:
        name = it.get('name', '')
        qty = int(it.get('quantity', 0))
        if item_type == 'oil' and name.startswith('Oil Change (') and name.endswith(')'):
            inner = name.replace('Oil Change (', '').replace(')', '')
            if inner == item_name:
                pending += qty
        elif item_type == 'wheel' and name.startswith('Wheel Change (') and name.endswith(')'):
            inner = name.replace('Wheel Change (', '').replace(')', '')
            if inner == item_name:
                pending += qty
        elif item_type == 'product' and not ('Oil Change (' in name or 'Wheel Change (' in name):
            if name == item_name:
                pending += qty
    return pending

def get_available_stock(table_name, item_name, item_type='product'):
    """Calculate available stock considering pending items"""
    base = get_stock(table_name, item_name)
    pending = pending_in_session(item_name, item_type)
    return base - pending

def update_stock(table_name, product_name, quantity):
    """Update stock after sale (subtract quantity)"""
    conn = get_db_connection()
    cur = conn.cursor()
    
    query = sql.SQL("UPDATE {} SET stock = stock - %s WHERE name = %s").format(
        sql.Identifier(table_name))
    
    cur.execute(query, (quantity, product_name))
    conn.commit()
    cur.close()
    conn.close()

# --------------------------
# Transaction Processing
# --------------------------
def is_service_or_used(item_name):
    """Return True if item is service or used part"""
    if item_name is None:
        return False
    name = str(item_name)
    return (name.startswith('Service') or name.startswith('Used Part') or 
            name.startswith('قطعة مستعملة') or name.startswith('Used Part:'))

def log_sale(receipt_items, receipt_id):
    """Record cash sale"""
    conn = get_db_connection()
    cur = conn.cursor()
    now = datetime.now()
    
    for item in receipt_items:
        product, price, qty = item['name'], float(item['price']), int(item['quantity'])
        total = price * qty
        
        cur.execute(
            "INSERT INTO sales (product, price, quantity, total, sale_date, receipt_id) "
            "VALUES (%s, %s, %s, %s, %s, %s)",
            (product, price, qty, total, now, receipt_id)
        )
        
        # Only update stock if it's a regular product (not service or used part)
        if not is_service_or_used(product):
            if product.startswith('Oil Change (') and product.endswith(')'):
                oil = product.replace('Oil Change (', '').replace(')', '')
                update_stock('oils', oil, qty)
            elif product.startswith('Wheel Change (') and product.endswith(')'):
                wheel = product.replace('Wheel Change (', '').replace(')', '')
                update_stock('wheels', wheel, qty)
            else:
                update_stock('products', product, qty)
    
    conn.commit()
    cur.close()
    conn.close()

def log_credit(receipt_items, receipt_id, customer_name):
    """Record credit sale"""
    conn = get_db_connection()
    cur = conn.cursor()
    now = datetime.now()
    month_year = now.strftime("%Y-%m")
    
    for item in receipt_items:
        product, price, qty = item['name'], float(item['price']), int(item['quantity'])
        total = price * qty
        
        cur.execute(
            "INSERT INTO credit_sales (customer_name, product, price, quantity, total, sale_date, receipt_id, month_year) "
            "VALUES (%s, %s, %s, %s, %s, %s, %s, %s)",
            (customer_name, product, price, qty, total, now, receipt_id, month_year)
        )
        
        if not is_service_or_used(product):
            if product.startswith('Oil Change (') and product.endswith(')'):
                oil = product.replace('Oil Change (', '').replace(')', '')
                update_stock('oils', oil, qty)
            elif product.startswith('Wheel Change (') and product.endswith(')'):
                wheel = product.replace('Wheel Change (', '').replace(')', '')
                update_stock('wheels', wheel, qty)
            else:
                update_stock('products', product, qty)
    
    conn.commit()
    cur.close()
    conn.close()

def log_medgulf(receipt_items, receipt_id, customer_name):
    """Record MedGulf sale"""
    conn = get_db_connection()
    cur = conn.cursor()
    now = datetime.now()
    month_year = now.strftime("%Y-%m")
    
    for item in receipt_items:
        product, price, qty = item['name'], float(item['price']), int(item['quantity'])
        total = price * qty
        
        cur.execute(
            "INSERT INTO medgulf_sales (customer_name, product, price, quantity, total, sale_date, receipt_id, month_year) "
            "VALUES (%s, %s, %s, %s, %s, %s, %s, %s)",
            (customer_name, product, price, qty, total, now, receipt_id, month_year)
        )
        
        if not is_service_or_used(product):
            if product.startswith('Oil Change (') and product.endswith(')'):
                oil = product.replace('Oil Change (', '').replace(')', '')
                update_stock('oils', oil, qty)
            elif product.startswith('Wheel Change (') and product.endswith(')'):
                wheel = product.replace('Wheel Change (', '').replace(')', '')
                update_stock('wheels', wheel, qty)
            else:
                update_stock('products', product, qty)
    
    conn.commit()
    cur.close()
    conn.close()

# --------------------------
# Report Generation
# --------------------------
def generate_daily_word_report():
    """Generate daily sales report"""
    today = datetime.now().strftime("%Y-%m-%d")
    doc = Document()
    doc.add_heading(f"Salimco Motorcycle Shop - Daily Report - {today}", level=1)

    # Normal sales
    doc.add_heading("=== Normal Sales ===", level=2)
    conn = get_db_connection()
    cur = conn.cursor()
    
    cur.execute(
        "SELECT product, price, quantity, total, sale_date, receipt_id "
        "FROM sales WHERE DATE(sale_date) = %s ORDER BY sale_date",
        (today,)
    )
    sales = cur.fetchall()
    
    if sales:
        tbl = doc.add_table(rows=1, cols=6)
        hdr = tbl.rows[0].cells
        hdr[0].text='Product'; hdr[1].text='Price'; hdr[2].text='Qty'; hdr[3].text='Total'; hdr[4].text='DateTime'; hdr[5].text='ReceiptID'
        subtotal = 0.0
        for r in sales:
            row = tbl.add_row().cells
            row[0].text = str(r[0])
            row[1].text = f"{float(r[1]):.2f}"
            row[2].text = str(int(r[2]))
            row[3].text = f"{float(r[3]):.2f}"
            row[4].text = str(r[4])
            row[5].text = str(r[5])
            subtotal += float(r[3])
        doc.add_paragraph(f"Subtotal (Normal): {subtotal:.2f}")
    else:
        doc.add_paragraph("No normal sales for today.")

    # Debts
    doc.add_heading("=== Debt Transactions ===", level=2)
    cur.execute(
        "SELECT customer_name, product, price, quantity, total, sale_date, receipt_id "
        "FROM credit_sales WHERE DATE(sale_date) = %s ORDER BY sale_date",
        (today,)
    )
    debts = cur.fetchall()
    
    if debts:
        tbl = doc.add_table(rows=1, cols=7)
        hdr = tbl.rows[0].cells
        hdr[0].text='Customer'; hdr[1].text='Product'; hdr[2].text='Price'; hdr[3].text='Qty'; hdr[4].text='Total'; hdr[5].text='DateTime'; hdr[6].text='ReceiptID'
        subtotal = 0.0
        for r in debts:
            row = tbl.add_row().cells
            row[0].text = str(r[0]); row[1].text = str(r[1])
            row[2].text = f"{float(r[2]):.2f}"
            row[3].text = str(int(r[3]))
            row[4].text = f"{float(r[4]):.2f}"
            row[5].text = str(r[5]); row[6].text = str(r[6])
            subtotal += float(r[4])
        doc.add_paragraph(f"Subtotal (Debts): {subtotal:.2f}")
    else:
        doc.add_paragraph("No debt transactions for today.")

    # MedGulf
    doc.add_heading("=== MedGulf Transactions ===", level=2)
    cur.execute(
        "SELECT customer_name, product, price, quantity, total, sale_date, receipt_id "
        "FROM medgulf_sales WHERE DATE(sale_date) = %s ORDER BY sale_date",
        (today,)
    )
    med = cur.fetchall()
    
    if med:
        tbl = doc.add_table(rows=1, cols=7)
        hdr = tbl.rows[0].cells
        hdr[0].text='Customer'; hdr[1].text='Product'; hdr[2].text='Price'; hdr[3].text='Qty'; hdr[4].text='Total'; hdr[5].text='DateTime'; hdr[6].text='ReceiptID'
        subtotal = 0.0
        for r in med:
            row = tbl.add_row().cells
            row[0].text = str(r[0]); row[1].text = str(r[1])
            row[2].text = f"{float(r[2]):.2f}"
            row[3].text = str(int(r[3]))
            row[4].text = f"{float(r[4]):.2f}"
            row[5].text = str(r[5]); row[6].text = str(r[6])
            subtotal += float(r[4])
        doc.add_paragraph(f"Subtotal (MedGulf): {subtotal:.2f}")
    else:
        doc.add_paragraph("No MedGulf transactions for today.")

    # Services & Used Parts
    doc.add_heading("=== Services and Used Parts Summary ===", level=2)
    cur.execute(
        "SELECT product, price, quantity, total, sale_date, receipt_id "
        "FROM sales WHERE DATE(sale_date) = %s AND "
        "(product LIKE 'Service%' OR product LIKE 'Used Part%' OR product LIKE 'قطعة مستعملة%') "
        "ORDER BY product",
        (today,)
    )
    service_items = cur.fetchall()
    
    if service_items:
        services = []
        used_parts = []
        for r in service_items:
            pname = str(r[0])
            if pname.lower().startswith('service'):
                services.append(r)
            else:
                used_parts.append(r)
        
        if services:
            tbl = doc.add_table(rows=1, cols=6)
            hdr = tbl.rows[0].cells
            hdr[0].text='Product'; hdr[1].text='Price'; hdr[2].text='Qty'; hdr[3].text='Total'; hdr[4].text='DateTime'; hdr[5].text='ReceiptID'
            for r in services:
                row = tbl.add_row().cells
                row[0].text = str(r[0])
                row[1].text = f"{float(r[1]):.2f}"
                row[2].text = str(int(r[2]))
                row[3].text = f"{float(r[3]):.2f}"
                row[4].text = str(r[4]); row[5].text = str(r[5])
        else:
            doc.add_paragraph("No services for today.")

        if used_parts:
            tbl = doc.add_table(rows=1, cols=6)
            hdr = tbl.rows[0].cells
            hdr[0].text='Product'; hdr[1].text='Price'; hdr[2].text='Qty'; hdr[3].text='Total'; hdr[4].text='DateTime'; hdr[5].text='ReceiptID'
            for r in used_parts:
                row = tbl.add_row().cells
                row[0].text = str(r[0])
                row[1].text = f"{float(r[1]):.2f}"
                row[2].text = str(int(r[2]))
                row[3].text = f"{float(r[3]):.2f}"
                row[4].text = str(r[4]); row[5].text = str(r[5])
        else:
            doc.add_paragraph("No used parts for today.")
    else:
        doc.add_paragraph("No services or used parts for today.")
    
    cur.close()
    conn.close()

    # Save the report
    os.makedirs('reports', exist_ok=True)
    filename = f"Daily_Report_{today}.docx"
    filepath = os.path.join('reports', filename)
    doc.save(filepath)
    return filepath

def generate_debts_word_report():
    """Generate detailed monthly debts report with customer subtotals"""
    doc = Document()
    month = datetime.now().strftime("%Y-%m")
    doc.add_heading(f"Salimco - Monthly Debts Report - {month}", level=1)
    
    conn = get_db_connection()
    cur = conn.cursor()
    
    # Get all transactions for the month
    cur.execute(
        "SELECT customer_name, product, price, quantity, total, sale_date, receipt_id "
        "FROM credit_sales WHERE month_year = %s ORDER BY customer_name, sale_date",
        (month,)
    )
    transactions = cur.fetchall()
    
    if not transactions:
        doc.add_paragraph("No debt transactions for this month.")
        cur.close()
        conn.close()
        
        os.makedirs('reports', exist_ok=True)
        filename = f"Debts_Report_{month}.docx"
        filepath = os.path.join('reports', filename)
        doc.save(filepath)
        return filepath
    
    # Calculate customer totals
    cur.execute(
        "SELECT customer_name, SUM(total) "
        "FROM credit_sales WHERE month_year = %s "
        "GROUP BY customer_name ORDER BY customer_name",
        (month,)
    )
    customer_totals = cur.fetchall()
    
    # Add customer summary
    doc.add_heading("Customer Debt Summary", level=2)
    summary_table = doc.add_table(rows=1, cols=2)
    hdr = summary_table.rows[0].cells
    hdr[0].text = 'Customer'
    hdr[1].text = 'Total Owed'
    
    grand_total = 0.0
    for customer, total in customer_totals:
        row = summary_table.add_row().cells
        row[0].text = customer
        row[1].text = f"{float(total):.2f}"
        grand_total += float(total)
    
    doc.add_paragraph(f"\nGrand Total: {grand_total:.2f}")
    
    # Add transaction details
    doc.add_heading("Transaction Details", level=2)
    details_table = doc.add_table(rows=1, cols=7)
    hdr = details_table.rows[0].cells
    hdr[0].text = 'Customer'
    hdr[1].text = 'Product'
    hdr[2].text = 'Price'
    hdr[3].text = 'Qty'
    hdr[4].text = 'Total'
    hdr[5].text = 'Date'
    hdr[6].text = 'Receipt'
    
    for t in transactions:
        row = details_table.add_row().cells
        row[0].text = str(t[0])
        row[1].text = str(t[1])
        row[2].text = f"{float(t[2]):.2f}"
        row[3].text = str(int(t[3]))
        row[4].text = f"{float(t[4]):.2f}"
        row[5].text = str(t[5])
        row[6].text = str(t[6])
    
    cur.close()
    conn.close()
    
    os.makedirs('reports', exist_ok=True)
    filename = f"Debts_Report_{month}.docx"
    filepath = os.path.join('reports', filename)
    doc.save(filepath)
    return filepath

def generate_medgulf_word_report():
    """Generate monthly MedGulf report"""
    month = datetime.now().strftime("%Y-%m")
    doc = Document()
    doc.add_heading(f"Salimco - MedGulf Report - {month}", level=1)
    
    conn = get_db_connection()
    cur = conn.cursor()
    
    cur.execute(
        "SELECT customer_name, product, price, quantity, total, sale_date, receipt_id "
        "FROM medgulf_sales WHERE month_year = %s ORDER BY sale_date",
        (month,)
    )
    med = cur.fetchall()
    
    if med:
        tbl = doc.add_table(rows=1, cols=7)
        hdr = tbl.rows[0].cells
        hdr[0].text='Customer'; hdr[1].text='Product'; hdr[2].text='Price'; hdr[3].text='Qty'; hdr[4].text='Total'; hdr[5].text='DateTime'; hdr[6].text='ReceiptID'
        subtotal = 0.0
        for r in med:
            row = tbl.add_row().cells
            row[0].text = str(r[0]); row[1].text = str(r[1])
            row[2].text = f"{float(r[2]):.2f}"
            row[3].text = str(int(r[3]))
            row[4].text = f"{float(r[4]):.2f}"
            row[5].text = str(r[5]); row[6].text = str(r[6])
            subtotal += float(r[4])
        doc.add_paragraph(f"Subtotal (MedGulf): {subtotal:.2f}")
    else:
        doc.add_paragraph("No MedGulf transactions for this month.")
    
    cur.close()
    conn.close()
    
    os.makedirs('reports', exist_ok=True)
    filename = f"MedGulf_Report_{month}.docx"
    filepath = os.path.join('reports', filename)
    doc.save(filepath)
    return filepath

# --------------------------
# Application Routes
# --------------------------
@app.route('/', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form.get('username','')
        password = request.form.get('password','')
        role = validate_login(username, password)
        if role:
            session['username'] = username
            session['role'] = role
            return redirect(url_for('pos'))
        flash('Invalid username or password', 'danger')
    return render_template('login.html', shop_name='Salimco Motorcycle Shop')

@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('login'))

@app.route('/pos', methods=['GET', 'POST'])
def pos():
    if 'username' not in session:
        return redirect(url_for('login'))

    if 'receipt_items' not in session:
        session['receipt_items'] = []
        session['receipt_id'] = datetime.now().strftime("%Y%m%d%H%M%S")

    if request.method == 'POST':
        action = request.form.get('action')

        if action == 'add_product':
            product_name = request.form.get('product_name')
            try:
                qty = int(request.form.get('quantity',1))
            except:
                qty = 1
            product = next((p for p in get_products() if p['name'] == product_name), None)
            if product:
                available = get_available_stock('products', product_name, 'product')
                if qty > available:
                    flash(f"Only {available} available in stock", 'warning')
                else:
                    session['receipt_items'].append({'name': product_name, 'price': product['price'], 'quantity': qty})
                    session.modified = True
                    flash('Added product to receipt', 'success')

        elif action == 'add_oil':
            oil_name = request.form.get('oil_name')
            try:
                qty = int(request.form.get('quantity',1))
            except:
                qty = 1
            oil = next((o for o in get_oils() if o['name'] == oil_name), None)
            if oil:
                available = get_available_stock('oils', oil_name, 'oil')
                if qty > available:
                    flash(f"Only {available} oil in stock", 'warning')
                else:
                    item_name = f"Oil Change ({oil_name})"
                    session['receipt_items'].append({'name': item_name, 'price': oil['price'], 'quantity': qty})
                    session.modified = True
                    flash('Added oil change to receipt', 'success')

        elif action == 'add_wheel':
            wheel_name = request.form.get('wheel_name')
            try:
                qty = int(request.form.get('quantity',1))
            except:
                qty = 1
            wheel = next((w for w in get_wheels() if w['name'] == wheel_name), None)
            if wheel:
                available = get_available_stock('wheels', wheel_name, 'wheel')
                if qty > available:
                    flash(f"Only {available} wheel in stock", 'warning')
                else:
                    item_name = f"Wheel Change ({wheel_name})"
                    session['receipt_items'].append({'name': item_name, 'price': wheel['price'], 'quantity': qty})
                    session.modified = True
                    flash('Added wheel change to receipt', 'success')

        elif action == 'add_service':
            service_name = request.form.get('service_name', 'Service Charge').strip()
            try:
                price = float(request.form.get('service_price', 0))
                if price <= 0:
                    flash('Service price must be positive', 'warning')
                else:
                    session['receipt_items'].append({'name': f"Service: {service_name}", 'price': price, 'quantity': 1})
                    session.modified = True
                    flash('Added service charge to receipt', 'success')
            except ValueError:
                flash('Invalid service price', 'danger')

        elif action == 'add_used_part':
            part_name = request.form.get('part_name', 'قطعة مستعملة').strip()
            try:
                price = float(request.form.get('part_price', 0))
                if price <= 0:
                    flash('Part price must be positive', 'warning')
                else:
                    session['receipt_items'].append({'name': f"Used Part: {part_name}", 'price': price, 'quantity': 1})
                    session.modified = True
                    flash('Added used part to receipt', 'success')
            except ValueError:
                flash('Invalid part price', 'danger')

        elif action == 'finalize_cash':
            if not session.get('receipt_items'):
                flash('Receipt empty', 'warning')
            else:
                rid = session.get('receipt_id')
                log_sale(session['receipt_items'], rid)
                session['receipt_items'] = []
                session['receipt_id'] = datetime.now().strftime("%Y%m%d%H%M%S")
                session.modified = True
                flash(f"Saved Receipt #{rid} (Cash)", 'success')

        elif action == 'finalize_credit':
            if not session.get('receipt_items'):
                flash('Receipt empty', 'warning')
            else:
                customer_name = request.form.get('customer_name','').strip()
                if not customer_name:
                    flash('Customer name required for credit', 'warning')
                else:
                    rid = session.get('receipt_id')
                    log_credit(session['receipt_items'], rid, customer_name)
                    session['receipt_items'] = []
                    session['receipt_id'] = datetime.now().strftime("%Y%m%d%H%M%S")
                    session.modified = True
                    flash(f"Saved Receipt #{rid} (Credit)", 'success')

        elif action == 'finalize_medgulf':
            if not session.get('receipt_items'):
                flash('Receipt empty', 'warning')
            else:
                customer_name = request.form.get('customer_name','').strip()
                if not customer_name:
                    flash('Customer name required for MedGulf', 'warning')
                else:
                    rid = session.get('receipt_id')
                    log_medgulf(session['receipt_items'], rid, customer_name)
                    session['receipt_items'] = []
                    session['receipt_id'] = datetime.now().strftime("%Y%m%d%H%M%S")
                    session.modified = True
                    flash(f"Saved Receipt #{rid} (MedGulf)", 'success')

        return redirect(url_for('pos'))

    total = sum(float(i['price']) * int(i['quantity']) for i in session.get('receipt_items', []))
    return render_template('pos.html',
                           products=get_products(),
                           oils=get_oils(),
                           wheels=get_wheels(),
                           receipt_items=session.get('receipt_items', []),
                           total=total,
                           receipt_id=session.get('receipt_id'),
                           role=session.get('role'),
                           shop_name='Salimco Motorcycle Shop')

@app.route('/remove_from_cart/<int:index>', methods=['POST'])
def remove_from_cart(index):
    items = session.get('receipt_items', [])
    if 0 <= index < len(items):
        removed = items.pop(index)
        session['receipt_items'] = items
        session.modified = True
        flash(f"Removed {removed['name']} x{removed['quantity']}", 'info')
    else:
        flash('Invalid item index', 'danger')
    return redirect(url_for('pos'))

@app.route('/inventory', methods=['GET', 'POST'])
def inventory():
    if 'username' not in session or session.get('role') != 'admin':
        flash('Access denied', 'danger')
        return redirect(url_for('pos'))

    q = request.args.get('q', '').strip().lower()

    if request.method == 'POST':
        action = request.form.get('action')
        if action == 'add_product':
            name = request.form.get('product_name','').strip()
            try:
                buy_price = float(request.form.get('buy_price',0))
            except:
                buy_price = 0.0
            try:
                sell_price = float(request.form.get('sell_price',0))
            except:
                sell_price = 0.0
            try:
                stock = int(request.form.get('stock',0))
            except:
                stock = 0
            
            conn = get_db_connection()
            cur = conn.cursor()
            try:
                cur.execute(
                    "INSERT INTO products (name, buy_price, sell_price, stock) "
                    "VALUES (%s, %s, %s, %s)",
                    (name, buy_price, sell_price, stock)
                )
                conn.commit()
                flash(f"Product '{name}' added", 'success')
            except psycopg2.IntegrityError:
                conn.rollback()
                flash(f"Product '{name}' already exists", 'danger')
            finally:
                cur.close()
                conn.close()

        elif action == 'add_oil':
            name = request.form.get('oil_name','').strip()
            try:
                buy_price = float(request.form.get('buy_price',0))
            except:
                buy_price = 0.0
            try:
                sell_price = float(request.form.get('sell_price',0))
            except:
                sell_price = 0.0
            try:
                stock = int(request.form.get('stock',0))
            except:
                stock = 0
            
            conn = get_db_connection()
            cur = conn.cursor()
            try:
                cur.execute(
                    "INSERT INTO oils (name, buy_price, sell_price, stock) "
                    "VALUES (%s, %s, %s, %s)",
                    (name, buy_price, sell_price, stock)
                )
                conn.commit()
                flash(f"Oil '{name}' added", 'success')
            except psycopg2.IntegrityError:
                conn.rollback()
                flash(f"Oil '{name}' already exists", 'danger')
            finally:
                cur.close()
                conn.close()

        elif action == 'add_wheel':
            name = request.form.get('wheel_name','').strip()
            try:
                buy_price = float(request.form.get('buy_price',0))
            except:
                buy_price = 0.0
            try:
                sell_price = float(request.form.get('sell_price',0))
            except:
                sell_price = 0.0
            try:
                stock = int(request.form.get('stock',0))
            except:
                stock = 0
            
            conn = get_db_connection()
            cur = conn.cursor()
            try:
                cur.execute(
                    "INSERT INTO wheels (name, buy_price, sell_price, stock) "
                    "VALUES (%s, %s, %s, %s)",
                    (name, buy_price, sell_price, stock)
                )
                conn.commit()
                flash(f"Wheel '{name}' added", 'success')
            except psycopg2.IntegrityError:
                conn.rollback()
                flash(f"Wheel '{name}' already exists", 'danger')
            finally:
                cur.close()
                conn.close()

        return redirect(url_for('inventory'))

    products = get_products()
    oils = get_oils()
    wheels = get_wheels()

    if q:
        products = [p for p in products if q in p['name'].lower()]
        oils = [o for o in oils if q in o['name'].lower()]
        wheels = [w for w in wheels if q in w['name'].lower()]

    return render_template('inventory.html',
                           products=products,
                           oils=oils,
                           wheels=wheels,
                           q=q,
                           shop_name='Salimco Motorcycle Shop')

@app.route('/report/daily')
def report_daily():
    try:
        filename = generate_daily_word_report()
        return send_file(filename, as_attachment=True)
    except Exception as e:
        flash(f"Error generating daily report: {str(e)}", "danger")
        return redirect(url_for('pos'))

@app.route('/report/debts')
def report_debts():
    try:
        filename = generate_debts_word_report()
        return send_file(filename, as_attachment=True)
    except Exception as e:
        flash(f"Error generating debts report: {str(e)}", "danger")
        return redirect(url_for('pos'))

@app.route('/report/medgulf')
def report_medgulf():
    try:
        filename = generate_medgulf_word_report()
        return send_file(filename, as_attachment=True)
    except Exception as e:
        flash(f"Error generating MedGulf report: {str(e)}", "danger")
        return redirect(url_for('pos'))

if __name__ == '__main__':
    initialize_database()
    app.run(host='0.0.0.0', port=5000, debug=False)